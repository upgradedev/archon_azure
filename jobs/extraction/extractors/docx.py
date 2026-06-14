"""DOCX / DOC extractor — uses python-docx then delegates to text LLM."""

import json
from pathlib import Path

from docx import Document
from tenacity import retry, stop_after_attempt, wait_exponential

from .base import BaseExtractor
from .image import EXTRACTION_PROMPT
from .utils import build_document, _clean_json
from models.document import ExtractedDocument


class DocxExtractor(BaseExtractor):
    def __init__(self):
        from services.llm import get_client
        self.client = get_client()
        self.model_name = __import__("os").getenv("AZURE_OPENAI_VISION_DEPLOYMENT", "gpt-4o")

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in {".docx", ".doc"}

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(min=2, max=10))
    def extract(self, path: Path) -> ExtractedDocument:
        text = _extract_docx_text(path)
        prompt = (
            "You are a financial document extraction specialist.\n"
            "Extract from the following document text (may be Greek or English).\n\n"
            f"DOCUMENT TEXT:\n{text[:4000]}\n\n"
            + EXTRACTION_PROMPT
        )
        response = self.client.chat.completions.create(
            model=self.model_name,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2048,
            temperature=0.1,
        )
        raw = response.choices[0].message.content or "{}"
        data = json.loads(_clean_json(raw.strip()))

        return build_document(
            source_file=path.name,
            data=data,
            extraction_model=self.model_name,
            raw_text_excerpt=text[:500],
            default_confidence=0.88,
        )


def _extract_docx_text(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)
